"""
Document parser module for SpecGuard.

Handles parsing of various document formats:
- PDF (via PyMuPDF/fitz)
- DOCX (via python-docx)
- Markdown (plain text)
- Plain text

Produces chunked, indexed text for downstream processing.
"""

import logging
import re
from pathlib import Path
from typing import Optional

from .models import DocumentChunk, ParsedDocument, SourceLocation

logger = logging.getLogger(__name__)

# Chunk configuration
DEFAULT_CHUNK_SIZE = 1500  # characters
DEFAULT_CHUNK_OVERLAP = 200  # characters


def parse_document(
    file_path: Path,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
) -> ParsedDocument:
    """
    Parse a document and return structured chunks.

    Args:
        file_path: Path to the document file
        chunk_size: Target size of each chunk in characters
        chunk_overlap: Overlap between consecutive chunks

    Returns:
        ParsedDocument with full text and chunks

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file does not exist
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    suffix = file_path.suffix.lower()
    filename = file_path.name

    logger.info(f"Parsing document: {filename} (type: {suffix})")

    if suffix == ".pdf":
        full_text, total_pages = _parse_pdf(file_path)
    elif suffix == ".docx":
        full_text, total_pages = _parse_docx(file_path)
    elif suffix in (".md", ".txt", ".text"):
        full_text, total_pages = _parse_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    chunks = _create_chunks(
        text=full_text,
        filename=filename,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )

    logger.info(f"Parsed {filename}: {len(full_text)} chars, {len(chunks)} chunks")

    return ParsedDocument(
        filename=filename,
        file_type=suffix,
        total_pages=total_pages,
        full_text=full_text,
        chunks=chunks
    )


def _parse_pdf(file_path: Path) -> tuple[str, int]:
    """
    Parse PDF using PyMuPDF.

    Returns:
        Tuple of (full_text, page_count)
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. "
            "Install with: pip install pymupdf"
        )

    doc = fitz.open(file_path)
    pages_text = []

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        # Add page marker for reference
        pages_text.append(f"[PAGE {page_num}]\n{text}")

    doc.close()

    full_text = "\n\n".join(pages_text)
    return full_text, len(pages_text)


def _parse_docx(file_path: Path) -> tuple[str, Optional[int]]:
    """
    Parse DOCX using python-docx.

    Returns:
        Tuple of (full_text, None) - DOCX doesn't have fixed pages
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
                paragraphs.append(f"{'#' * int(level)} {text}")
            else:
                paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text.append(row_text)
        if table_text:
            paragraphs.append("\n".join(table_text))

    full_text = "\n\n".join(paragraphs)
    return full_text, None


def _parse_text(file_path: Path) -> tuple[str, Optional[int]]:
    """
    Parse plain text or markdown file.

    Returns:
        Tuple of (full_text, None)
    """
    with open(file_path, "r", encoding="utf-8") as f:
        full_text = f.read()

    return full_text, None


def _create_chunks(
    text: str,
    filename: str,
    chunk_size: int,
    chunk_overlap: int
) -> list[DocumentChunk]:
    """
    Split text into overlapping chunks with location metadata.

    Uses paragraph boundaries when possible for more coherent chunks.
    """
    chunks = []

    # Split by paragraphs first
    paragraphs = re.split(r'\n\s*\n', text)

    current_chunk = ""
    current_start_line = 1
    line_count = 1
    chunk_index = 0

    for para in paragraphs:
        para_lines = para.count('\n') + 1

        # If adding this paragraph exceeds chunk size, save current and start new
        if len(current_chunk) + len(para) > chunk_size and current_chunk:
            # Detect section/page from content
            section = _extract_section(current_chunk)
            page = _extract_page(current_chunk)

            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                source_location=SourceLocation(
                    filename=filename,
                    page=page,
                    section=section,
                    line_start=current_start_line,
                    line_end=line_count - 1
                ),
                chunk_index=chunk_index,
                keywords=_extract_keywords(current_chunk)
            ))
            chunk_index += 1

            # Start new chunk with overlap
            overlap_text = current_chunk[-chunk_overlap:] if len(current_chunk) > chunk_overlap else current_chunk
            current_chunk = overlap_text + "\n\n" + para
            current_start_line = max(1, line_count - overlap_text.count('\n'))
        else:
            if current_chunk:
                current_chunk += "\n\n" + para
            else:
                current_chunk = para

        line_count += para_lines

    # Don't forget the last chunk
    if current_chunk.strip():
        section = _extract_section(current_chunk)
        page = _extract_page(current_chunk)

        chunks.append(DocumentChunk(
            content=current_chunk.strip(),
            source_location=SourceLocation(
                filename=filename,
                page=page,
                section=section,
                line_start=current_start_line,
                line_end=line_count
            ),
            chunk_index=chunk_index,
            keywords=_extract_keywords(current_chunk)
        ))

    return chunks


def _extract_section(text: str) -> Optional[str]:
    """Extract section number from text if present."""
    # Look for patterns like "Section 5.2.1" or "5.2.1" at start of text
    patterns = [
        r'(?:Section|§)\s*(\d+(?:\.\d+)*)',  # Section 5.2.1 or § 5.2.1
        r'^#+\s*(\d+(?:\.\d+)*)',  # Markdown heading with number
        r'^(\d+(?:\.\d+)+)\s+\w',  # 5.2.1 Title
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            return match.group(1)

    return None


def _extract_page(text: str) -> Optional[int]:
    """Extract page number from text if present (from PDF markers)."""
    match = re.search(r'\[PAGE\s+(\d+)\]', text)
    if match:
        return int(match.group(1))
    return None


def _extract_keywords(text: str) -> list[str]:
    """
    Extract keywords from text for retrieval indexing.

    Focuses on technical terms, nouns, and important identifiers.
    """
    # Remove common stop words and extract significant terms
    stop_words = {
        'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
        'should', 'may', 'might', 'must', 'shall', 'can', 'need', 'to', 'of',
        'in', 'for', 'on', 'with', 'at', 'by', 'from', 'as', 'into', 'through',
        'during', 'before', 'after', 'above', 'below', 'between', 'under',
        'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where',
        'why', 'how', 'all', 'each', 'few', 'more', 'most', 'other', 'some',
        'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too',
        'very', 'just', 'and', 'but', 'if', 'or', 'because', 'until', 'while',
        'this', 'that', 'these', 'those', 'it', 'its', 'page'
    }

    # Extract words
    words = re.findall(r'\b[a-zA-Z_][a-zA-Z0-9_]*\b', text.lower())

    # Filter and deduplicate
    keywords = []
    seen = set()
    for word in words:
        if word not in stop_words and len(word) > 2 and word not in seen:
            keywords.append(word)
            seen.add(word)

    # Return top keywords by frequency indication (order of appearance)
    return keywords[:50]


def parse_from_bytes(
    content: bytes,
    filename: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
) -> ParsedDocument:
    """
    Parse document from bytes content (for file uploads).

    Args:
        content: Raw file bytes
        filename: Original filename (used to determine type)
        chunk_size: Target chunk size
        chunk_overlap: Overlap between chunks

    Returns:
        ParsedDocument with full text and chunks
    """
    import tempfile
    from pathlib import Path

    # Write to temp file for parsing
    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        result = parse_document(tmp_path, chunk_size, chunk_overlap)
        # Update filename to original
        result.filename = filename
        for chunk in result.chunks:
            chunk.source_location.filename = filename
        return result
    finally:
        tmp_path.unlink()
